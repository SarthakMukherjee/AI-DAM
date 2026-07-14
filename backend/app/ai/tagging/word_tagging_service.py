from docx import Document

class WordTaggingService:
    def extract_text(self, word_path: str) -> str:
        try:
            doc = Document(word_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            return "\n".join(full_text).strip()
        except Exception as e:
            print(f"Word text extraction failed: {e}")
            return ""

    def extract_metadata(self, word_path: str) -> dict:
        try:
            doc = Document(word_path)
            core_props = doc.core_properties
            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or ""
            }
        except Exception as e:
            print(f"Word metadata extraction failed: {e}")
            return {
                "title": "",
                "author": "",
                "subject": "",
                "keywords": ""
            }
